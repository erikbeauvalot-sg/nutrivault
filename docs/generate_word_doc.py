"""
Génère le guide utilisateur NutriVault en format Word (.docx)
avec mise en page professionnelle et espaces pour captures d'écran.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# ─── Palette couleurs NutriVault ─────────────────────────────────────────────
COLOR_PRIMARY   = RGBColor(0x2E, 0x86, 0xAB)   # bleu NutriVault
COLOR_ACCENT    = RGBColor(0xF7, 0xB7, 0x31)   # or solarpunk
COLOR_SUCCESS   = RGBColor(0x27, 0xAE, 0x60)   # vert
COLOR_LIGHT_BG  = RGBColor(0xF8, 0xF5, 0xEE)   # parchemin
COLOR_TEXT      = RGBColor(0x2C, 0x3E, 0x50)   # texte foncé
COLOR_MUTED     = RGBColor(0x7F, 0x8C, 0x8D)   # gris
COLOR_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_HEADER_BG = RGBColor(0x2E, 0x86, 0xAB)


def rgb_hex(color: RGBColor) -> str:
    return '{:02X}{:02X}{:02X}'.format(color[0], color[1], color[2])


def set_cell_bg(cell, color: RGBColor):
    """Applique une couleur de fond à une cellule de tableau."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(color))
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """Applique des bordures à une cellule."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing borders if any
    for existing in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(existing)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), kwargs.get('val', 'single'))
        border.set(qn('w:sz'), str(kwargs.get('sz', 4)))
        border.set(qn('w:color'), kwargs.get('color', '2E86AB'))
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_screenshot_placeholder(doc, label: str, height_cm: float = 7.0):
    """Ajoute un espace réservé pour une capture d'écran."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f'📸  {label}')
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_MUTED
    run.font.italic = True

    # Cadre visuel via un tableau 1x1
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Cm(15)
    # Hauteur via XML
    tr = cell._tc.getparent()
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))  # 567 twips/cm
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

    set_cell_bg(cell, RGBColor(0xF0, 0xF4, 0xF8))
    set_cell_border(cell, val='single', sz=6, color='CBD5E0')

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p.add_run('[ Insérer la capture d\'écran ici ]')
    run2.font.color.rgb = COLOR_MUTED
    run2.font.italic = True
    run2.font.size = Pt(10)

    doc.add_paragraph()  # espace après


def add_tip_box(doc, text: str, icon: str = '💡', color: RGBColor = None):
    """Ajoute une boîte d'astuce/avertissement colorée."""
    if color is None:
        color = RGBColor(0xFF, 0xF9, 0xE6)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, color)
    set_cell_border(cell, val='single', sz=8, color=rgb_hex(COLOR_ACCENT))
    p = cell.paragraphs[0]
    run = p.add_run(f'{icon}  {text}')
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_TEXT
    doc.add_paragraph()


def add_styled_table(doc, headers, rows, col_widths=None):
    """Ajoute un tableau stylé avec en-têtes colorés."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    # En-têtes
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, COLOR_PRIMARY)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = COLOR_WHITE
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Données
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = RGBColor(0xF8, 0xF9, 0xFA) if r_idx % 2 == 0 else COLOR_WHITE
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR_TEXT

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()


def add_section_header(doc, number: str, title: str):
    """Ajoute un titre de section avec barre de couleur via tableau."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Barre colorée
    bar_cell = tbl.cell(0, 0)
    bar_cell.width = Cm(0.4)
    set_cell_bg(bar_cell, COLOR_PRIMARY)
    bar_cell.paragraphs[0].add_run(' ')

    text_cell = tbl.cell(0, 1)
    set_cell_bg(text_cell, RGBColor(0xEC, 0xF3, 0xF9))
    p = text_cell.paragraphs[0]
    run_num = p.add_run(f'{number}  ')
    run_num.font.color.rgb = COLOR_PRIMARY
    run_num.font.bold = True
    run_num.font.size = Pt(14)
    run = p.add_run(title)
    run.font.color.rgb = COLOR_TEXT
    run.font.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()


def add_subsection(doc, title: str):
    h = doc.add_heading(title, level=3)
    h.runs[0].font.color.rgb = COLOR_PRIMARY
    h.runs[0].font.size = Pt(12)


def body_text(doc, text: str):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.color.rgb = COLOR_TEXT
    return p


def bullet(doc, text: str, level: int = 0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_TEXT
    return p


# ─── GÉNÉRATION DU DOCUMENT ──────────────────────────────────────────────────

doc = Document()

# Marges
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── PAGE DE COUVERTURE ───────────────────────────────────────────────────────
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.add_run('\n\n')

title_run = cover.add_run('NutriVault\n')
title_run.font.size = Pt(36)
title_run.font.bold = True
title_run.font.color.rgb = COLOR_PRIMARY

subtitle_run = cover.add_run('Guide d\'utilisation\nTemplates de notes de consultation\n\n')
subtitle_run.font.size = Pt(20)
subtitle_run.font.color.rgb = COLOR_ACCENT
subtitle_run.font.bold = True

version_run = cover.add_run('Version 8.7.16  ·  À l\'attention des diététicien(ne)s\n\n\n')
version_run.font.size = Pt(11)
version_run.font.color.rgb = COLOR_MUTED

doc.add_page_break()

# ── TABLE DES MATIÈRES (manuelle) ───────────────────────────────────────────
h = doc.add_heading('Table des matières', level=1)
h.runs[0].font.color.rgb = COLOR_PRIMARY

toc_items = [
    ('1.', 'Introduction', '3'),
    ('2.', 'Accéder aux templates', '3'),
    ('3.', 'Liste des templates', '3'),
    ('4.', 'Créer un nouveau template', '4'),
    ('    4.1', 'Informations générales', '4'),
    ('    4.2', 'Ajouter des éléments', '5'),
    ('    4.3', 'Les types d\'éléments', '5'),
    ('    4.4', 'Aperçu en temps réel', '7'),
    ('5.', 'Utiliser un template en consultation', '7'),
    ('    5.1', 'Démarrer une consultation', '7'),
    ('    5.2', 'Remplir la note', '8'),
    ('    5.3', 'Enregistrement automatique', '9'),
    ('    5.4', 'Finaliser et facturer', '9'),
    ('6.', 'Gérer ses templates', '10'),
    ('7.', 'Bonnes pratiques', '11'),
]

for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    tab_stop = p.paragraph_format.tab_stops
    run = p.add_run(f'{num}  {title}')
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_TEXT if not num.startswith('    ') else COLOR_MUTED

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
add_section_header(doc, '1', 'Introduction')

body_text(doc, (
    'Les templates de notes de consultation permettent de standardiser vos prises de notes '
    'lors des séances avec vos patients. Au lieu de recommencer une structure vierge à chaque '
    'consultation, vous préparez un modèle une seule fois — NutriVault l\'applique automatiquement.'
))

body_text(doc, 'Ce que vous pouvez inclure dans un template :')

add_styled_table(doc,
    headers=['Type d\'élément', 'Description'],
    rows=[
        ['Catégorie de champs', 'Regroupement de champs personnalisés (antécédents, habitudes alimentaires, etc.)'],
        ['Mesure clinique', 'Poids, IMC, tension artérielle, glycémie, etc.'],
        ['Instruction', 'Texte de guidage visible lors de la consultation, avec zone de notes libres'],
    ],
    col_widths=[5, 11]
)

body_text(doc, 'Avantages :')
bullet(doc, 'Gain de temps : structure préparée à l\'avance')
bullet(doc, 'Cohérence : mêmes informations collectées pour chaque type de consultation')
bullet(doc, 'Traçabilité : notes sauvegardées automatiquement, liées à la visite et au patient')
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. ACCÉDER AUX TEMPLATES
# ═══════════════════════════════════════════════════════════════════════════════
add_section_header(doc, '2', 'Accéder aux templates')

body_text(doc, 'Dans le menu de navigation latéral, cliquez sur "Templates de consultation".')
add_screenshot_placeholder(doc, 'Menu latéral — "Templates de consultation" mis en évidence', 5)

# ═══════════════════════════════════════════════════════════════════════════════
# 3. LISTE DES TEMPLATES
# ═══════════════════════════════════════════════════════════════════════════════
add_section_header(doc, '3', 'Liste des templates')

body_text(doc, 'La page d\'accueil des templates affiche tous vos modèles existants.')
add_screenshot_placeholder(doc, 'Page liste des templates', 7)

body_text(doc, 'Ce que vous voyez pour chaque template :')
bullet(doc, 'Nom avec son type (badge coloré)')
bullet(doc, 'Description courte')
bullet(doc, 'Nombre d\'éléments contenus')
bullet(doc, 'Visibilité : Privé ou Partagé avec l\'équipe')
bullet(doc, '⭐ Par défaut : si ce template est votre favori')
bullet(doc, 'Actions : Modifier · Dupliquer · Supprimer')
doc.add_paragraph()

body_text(doc, 'Types de templates disponibles :')
add_styled_table(doc,
    headers=['Type', 'Couleur', 'Usage conseillé'],
    rows=[
        ['Anamnèse', 'Rouge', 'Première consultation, collecte de l\'historique'],
        ['Évaluation', 'Bleu', 'Bilan nutritionnel, évaluation clinique'],
        ['Plan alimentaire', 'Vert', 'Consultation de prescription diététique'],
        ['Suivi', 'Orange', 'Consultations de suivi régulier'],
        ['Général', 'Violet', 'Usage polyvalent'],
        ['Personnalisé', 'Sarcelle', 'Tout autre cas spécifique'],
    ],
    col_widths=[4, 3, 9]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. CRÉER UN NOUVEAU TEMPLATE
# ═══════════════════════════════════════════════════════════════════════════════
add_section_header(doc, '4', 'Créer un nouveau template')

body_text(doc, 'Cliquez sur le bouton "+ Nouveau template" en haut à droite de la liste.')
add_screenshot_placeholder(doc, 'Bouton "+ Nouveau template"', 3.5)

body_text(doc, 'L\'éditeur de template s\'ouvre avec la configuration à gauche et l\'aperçu à droite.')

# 4.1
add_subsection(doc, '4.1  Informations générales')
add_screenshot_placeholder(doc, 'Section "Informations du template" — champs de configuration', 7)

body_text(doc, 'Renseignez les champs suivants :')
add_styled_table(doc,
    headers=['Champ', 'Obligatoire', 'Description'],
    rows=[
        ['Nom', '✅ Oui', 'Ex. : Bilan nutritionnel initial (200 caractères max)'],
        ['Type', 'Non', 'Catégorisation du template (Anamnèse, Évaluation, Suivi, etc.)'],
        ['Description', 'Non', 'Aide à choisir le bon template en consultation'],
        ['Visibilité', 'Non', 'Privé = vous seul(e) · Partagé = toute l\'équipe'],
        ['Couleur', 'Non', 'Code couleur pour identifier visuellement le template'],
        ['Template par défaut', 'Non', 'Ce template sera pré-sélectionné à chaque nouvelle consultation'],
        ['Tags', 'Non', 'Mots-clés : adulte, diabète, grossesse, sport, etc.'],
    ],
    col_widths=[4.5, 2.5, 9]
)

add_tip_box(doc,
    'Un seul template peut être marqué "par défaut". En activer un nouveau désactive automatiquement l\'ancien.',
    icon='💡'
)

# 4.2
add_subsection(doc, '4.2  Ajouter des éléments')
add_screenshot_placeholder(doc, 'Section "Éléments du template" avec plusieurs éléments ajoutés', 7)

body_text(doc, 'Pour ajouter un élément, cliquez sur l\'un des trois boutons disponibles :')
bullet(doc, '[ + Catégorie ]  — pour ajouter une catégorie de champs personnalisés')
bullet(doc, '[ + Mesure ]  — pour ajouter une mesure clinique')
bullet(doc, '[ + Instruction ]  — pour ajouter un bloc de texte de guidage')

add_screenshot_placeholder(doc, 'Boutons d\'ajout d\'éléments [ + Catégorie ] [ + Mesure ] [ + Instruction ]', 3.5)

body_text(doc, 'Pour réorganiser les éléments :')
bullet(doc, 'Utilisez les flèches ▲ ▼ sur chaque élément')
bullet(doc, 'Ou faites-les glisser via l\'icône ≡ à gauche (drag & drop)')

body_text(doc, 'Pour chaque élément, vous pouvez activer le toggle "Obligatoire" si ce champ doit impérativement être rempli avant de terminer la consultation.')
doc.add_paragraph()

doc.add_page_break()

# 4.3
add_subsection(doc, '4.3  Les types d\'éléments')

# Catégorie
h3 = doc.add_heading('▸  Catégorie de champs personnalisés', level=4)
h3.runs[0].font.color.rgb = COLOR_PRIMARY
h3.runs[0].font.size = Pt(11)

body_text(doc, 'Un clic sur "+ Catégorie" ouvre une fenêtre de sélection affichant toutes les catégories de champs disponibles.')
add_screenshot_placeholder(doc, 'Modal de sélection d\'une catégorie de champs', 7)

body_text(doc, 'Chaque catégorie affiche :')
bullet(doc, 'Nom et description')
bullet(doc, 'Nombre de champs inclus')
bullet(doc, 'Indicateur de niveau :')
bullet(doc, '🔵 Patient — données partagées entre toutes les consultations (ex. : antécédents médicaux)', level=1)
bullet(doc, '🟠 Par consultation — données spécifiques à cette séance (ex. : évaluation du jour)', level=1)

add_tip_box(doc,
    'La catégorie "Antécédents médicaux" est de niveau Patient — ses données s\'affichent dans '
    'toutes les consultations sans avoir à les re-saisir. La catégorie "Évaluation alimentaire du jour" '
    'est Par consultation — ses données concernent uniquement la séance en cours.',
    icon='💡'
)

add_screenshot_placeholder(doc, 'Élément catégorie dans la liste du template — badge de niveau visible', 4)

# Mesure
h3 = doc.add_heading('▸  Mesure clinique', level=4)
h3.runs[0].font.color.rgb = COLOR_PRIMARY
h3.runs[0].font.size = Pt(11)

body_text(doc, 'Un clic sur "+ Mesure" ouvre la liste des mesures cliniques disponibles dans votre compte (poids, IMC, tension artérielle, glycémie, tour de taille, etc.)')
add_screenshot_placeholder(doc, 'Modal de sélection d\'une mesure clinique', 7)

body_text(doc, 'Chaque mesure affiche son unité (kg, cm, mm Hg…) et sa catégorie.')
add_screenshot_placeholder(doc, 'Élément mesure dans la liste du template', 3.5)

doc.add_page_break()

# Instruction
h3 = doc.add_heading('▸  Instruction', level=4)
h3.runs[0].font.color.rgb = COLOR_PRIMARY
h3.runs[0].font.size = Pt(11)

body_text(doc, 'Une instruction est un bloc de texte que vous rédigez à l\'avance pour vous guider pendant la consultation. Elle s\'affiche en surligné lors de la prise de note et est accompagnée d\'une zone de notes libres.')
add_screenshot_placeholder(doc, 'Élément instruction en cours d\'édition dans le template', 5)

body_text(doc, 'Champs à remplir :')
bullet(doc, 'Titre  — Ex. : "Évaluation des habitudes alimentaires"')
bullet(doc, 'Contenu  — Ex. : "Demander au patient de décrire ses repas types de la semaine. Explorer les horaires, les lieux de repas, les préférences et les aversions."')

add_tip_box(doc,
    'Usage typique : guide de questionnement clinique, rappel de protocole, check-list d\'éléments à explorer.',
    icon='💡'
)

# 4.4
add_subsection(doc, '4.4  Aperçu en temps réel')

body_text(doc, 'Cliquez sur le bouton "Aperçu" (en haut à droite de l\'éditeur) pour afficher un panneau latéral montrant le rendu exact du template tel que vous le verrez lors d\'une consultation.')
add_screenshot_placeholder(doc, 'Éditeur avec panneau d\'aperçu ouvert à droite', 8)

body_text(doc, 'L\'aperçu se met à jour en temps réel au fur et à mesure que vous ajoutez ou modifiez les éléments.')

body_text(doc, 'Cliquez sur "Enregistrer" en haut à droite pour sauvegarder. Une confirmation verte apparaît.')
add_screenshot_placeholder(doc, 'Bouton Enregistrer et message de confirmation', 3.5)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. UTILISER UN TEMPLATE EN CONSULTATION
# ═══════════════════════════════════════════════════════════════════════════════
add_section_header(doc, '5', 'Utiliser un template en consultation')

# 5.1
add_subsection(doc, '5.1  Démarrer une consultation')

body_text(doc, 'Pour démarrer une consultation à partir d\'un template :')
bullet(doc, 'Ouvrez la fiche du patient')
bullet(doc, 'Cliquez sur la visite concernée pour accéder à son détail')
bullet(doc, 'Cliquez sur le bouton "Démarrer la consultation"')

add_screenshot_placeholder(doc, 'Bouton "Démarrer la consultation" dans la page de détail d\'une visite', 6)

body_text(doc, 'Une fenêtre s\'ouvre pour sélectionner un template.')
add_screenshot_placeholder(doc, 'Modal de sélection du template — cartes des templates disponibles', 7)

body_text(doc, 'Chaque carte affiche le nom, le type, la description et le nombre d\'éléments. Cliquez sur le template souhaité — NutriVault crée automatiquement la note et vous redirige vers l\'éditeur de consultation.')

# 5.2
add_subsection(doc, '5.2  Remplir la note')

add_screenshot_placeholder(doc, 'Page de l\'éditeur de note de consultation — vue d\'ensemble', 9)

body_text(doc, 'L\'éditeur affiche les éléments du template dans l\'ordre configuré. En-tête de la page : nom du template, statut (Brouillon / Complété), patient, date et diététicienne.')

doc.add_page_break()

h4 = doc.add_heading('Remplir une catégorie de champs', level=4)
h4.runs[0].font.color.rgb = COLOR_PRIMARY
h4.runs[0].font.size = Pt(11)
add_screenshot_placeholder(doc, 'Carte de catégorie avec champs à remplir', 7)
body_text(doc, 'Les champs apparaissent dans la mise en page configurée dans la catégorie (liste verticale ou grille). Remplissez normalement selon le type de champ (texte, liste déroulante, case à cocher, etc.)')

h4 = doc.add_heading('Saisir une mesure', level=4)
h4.runs[0].font.color.rgb = COLOR_PRIMARY
h4.runs[0].font.size = Pt(11)
add_screenshot_placeholder(doc, 'Carte de mesure clinique avec champ de saisie et unité', 5)
body_text(doc, 'Entrez la valeur numérique dans le champ prévu. L\'unité est affichée à côté.')

h4 = doc.add_heading('Utiliser une instruction', level=4)
h4.runs[0].font.color.rgb = COLOR_PRIMARY
h4.runs[0].font.size = Pt(11)
add_screenshot_placeholder(doc, 'Carte d\'instruction — encadré jaune de guidage et zone de notes', 6)
body_text(doc, 'Le contenu de l\'instruction s\'affiche dans un encadré jaune — c\'est votre aide-mémoire. La zone blanche en dessous est votre zone de notes libres pour documenter ce que le patient a répondu ou observé.')

h4 = doc.add_heading('Résumé de consultation', level=4)
h4.runs[0].font.color.rgb = COLOR_PRIMARY
h4.runs[0].font.size = Pt(11)
add_screenshot_placeholder(doc, 'Zone de résumé global en bas de page', 5)
body_text(doc, 'En bas de page, une grande zone texte permet de saisir un résumé global de la consultation.')

doc.add_page_break()

# 5.3
add_subsection(doc, '5.3  Enregistrement automatique')

add_screenshot_placeholder(doc, 'En-tête de la note avec indicateur "Sauvegarde auto" visible', 4)

body_text(doc, 'Vos notes sont sauvegardées automatiquement 3 secondes après chaque modification. L\'en-tête affiche :')
bullet(doc, '🔄 Spinner → sauvegarde en cours')
bullet(doc, '✅ "Sauvegarde auto" (vert) → sauvegarde réussie')
bullet(doc, '❌ "Échec de la sauvegarde auto" (rouge) → problème réseau, réessayez manuellement')

add_tip_box(doc,
    'Ne fermez pas l\'onglet si l\'indicateur "Sauvegarde en cours" est visible. '
    'Attendez la confirmation verte. Vous pouvez aussi cliquer sur "Enregistrer" à tout moment pour forcer une sauvegarde manuelle.',
    icon='⚠️',
    color=RGBColor(0xFF, 0xF0, 0xE6)
)

# 5.4
add_subsection(doc, '5.4  Finaliser et facturer')

body_text(doc, 'Quand la consultation est terminée, cliquez sur "Terminer & Facturer".')
add_screenshot_placeholder(doc, 'Bouton "Terminer & Facturer" dans l\'en-tête de la note', 3.5)

body_text(doc, 'Une fenêtre de confirmation apparaît avec trois options :')
add_screenshot_placeholder(doc, 'Modal "Terminer & Facturer" avec les trois options', 5)

add_styled_table(doc,
    headers=['Option', 'Description'],
    rows=[
        ['✅ Marquer la visite comme TERMINÉE', 'Change le statut de la visite dans NutriVault'],
        ['✅ Générer automatiquement une facture', 'Crée une facture associée à cette visite'],
        ['☐ Envoyer la facture par email', 'Envoie la facture au patient (si email renseigné)'],
    ],
    col_widths=[7, 9]
)

body_text(doc, 'Cliquez "Confirmer" — la note passe en statut Complétée et ne peut plus être modifiée.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. GÉRER SES TEMPLATES
# ═══════════════════════════════════════════════════════════════════════════════
add_section_header(doc, '6', 'Gérer ses templates')

add_subsection(doc, 'Dupliquer un template')
body_text(doc, 'Pour créer un template similaire à un existant, cliquez sur l\'icône Dupliquer dans la liste. Une copie identique est créée avec le suffixe (Copie) — vous pouvez ensuite la modifier librement.')
add_screenshot_placeholder(doc, 'Bouton Dupliquer dans la liste des templates', 4)

add_subsection(doc, 'Modifier un template existant')
body_text(doc, 'Cliquez sur Modifier (icône crayon). L\'éditeur s\'ouvre avec la configuration actuelle.')
add_tip_box(doc,
    'Modifier un template ne modifie PAS les notes déjà créées avec ce template. Seules les nouvelles consultations bénéficieront des changements.',
    icon='⚠️',
    color=RGBColor(0xFF, 0xF0, 0xE6)
)

add_subsection(doc, 'Supprimer un template')
body_text(doc, 'Cliquez sur l\'icône Supprimer (corbeille). Si des notes ont déjà été créées avec ce template, il sera désactivé (et non supprimé définitivement) pour préserver l\'historique.')

# ═══════════════════════════════════════════════════════════════════════════════
# 7. BONNES PRATIQUES
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_section_header(doc, '7', 'Bonnes pratiques')

add_subsection(doc, 'Organisez par type de consultation')
body_text(doc, 'Créez un template distinct pour chaque type de séance :')
bullet(doc, '"Bilan initial adulte" → type Anamnèse')
bullet(doc, '"Suivi mensuel" → type Suivi')
bullet(doc, '"Consultation diabète" → type Évaluation')
doc.add_paragraph()

add_subsection(doc, 'Utilisez les tags')
body_text(doc, 'Ajoutez des tags pour retrouver facilement vos templates : adulte, enfant, diabète, grossesse, sport, etc.')
doc.add_paragraph()

add_subsection(doc, 'Définissez un template par défaut')
body_text(doc, 'Si vous avez un type de consultation récurrent, marquez ce template comme "Par défaut" — il sera pré-sélectionné automatiquement à chaque nouvelle consultation.')
doc.add_paragraph()

add_subsection(doc, 'Partagez avec votre équipe')
body_text(doc, 'Si vous travaillez en cabinet avec d\'autres diététiciennes, passez vos templates en "Partagé" — toute l\'équipe peut les utiliser et vous évitez les doublons.')
doc.add_paragraph()

add_subsection(doc, 'Instructions comme aide-mémoire')
body_text(doc, 'Rédigez des instructions détaillées pour les protocoles complexes : liste de questions à poser, éléments à évaluer, grilles de scoring. Vos notes libres documenteront les réponses du patient directement en regard du protocole.')
doc.add_paragraph()

# ── PIED DE PAGE ─────────────────────────────────────────────────────────────
for section in doc.sections:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('NutriVault · Guide Templates de notes de consultation · Version 8.7.16')
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_MUTED

# ── SAUVEGARDE ────────────────────────────────────────────────────────────────
output_path = '/Users/erik/Documents/Dev/AI/Antigravity/nutrivault/docs/guide-templates-notes-consultation.docx'
doc.save(output_path)
print(f'✅ Document Word généré : {output_path}')
