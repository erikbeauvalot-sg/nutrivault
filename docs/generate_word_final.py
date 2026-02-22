"""
Génère le guide utilisateur NutriVault avec les vraies captures d'écran.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SCREENSHOTS = os.path.join(os.path.dirname(__file__), 'screenshots')
OUTPUT = os.path.join(os.path.dirname(__file__), 'guide-templates-notes-consultation.docx')

COLOR_PRIMARY  = RGBColor(0x2E, 0x86, 0xAB)
COLOR_ACCENT   = RGBColor(0xF7, 0xB7, 0x31)
COLOR_TEXT     = RGBColor(0x2C, 0x3E, 0x50)
COLOR_MUTED    = RGBColor(0x7F, 0x8C, 0x8D)
COLOR_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_TIP_BG   = RGBColor(0xFF, 0xF9, 0xE6)
COLOR_WARN_BG  = RGBColor(0xFF, 0xF0, 0xE6)

def rgb_hex(c): return '{:02X}{:02X}{:02X}'.format(c[0], c[1], c[2])

def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(color))
    tcPr.append(shd)

def set_cell_border(cell, color_hex='2E86AB', sz=6):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(existing)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:color'), color_hex)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def add_screenshot(doc, filename, caption=None, width_cm=15.5):
    filepath = os.path.join(SCREENSHOTS, filename)
    if os.path.exists(filepath):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(filepath, width=Cm(width_cm))
        if caption:
            cp = doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cp.add_run(f'↑ {caption}')
            cr.font.size = Pt(9)
            cr.font.italic = True
            cr.font.color.rgb = COLOR_MUTED
        doc.add_paragraph()
    else:
        # Placeholder si screenshot manquant
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f'[ {filename} — capture non disponible ]')
        r.font.color.rgb = COLOR_MUTED
        r.font.italic = True
        doc.add_paragraph()

def add_section_header(doc, number, title):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    bar = tbl.cell(0, 0)
    bar.width = Cm(0.4)
    set_cell_bg(bar, COLOR_PRIMARY)
    bar.paragraphs[0].add_run(' ')
    txt = tbl.cell(0, 1)
    set_cell_bg(txt, RGBColor(0xEC, 0xF3, 0xF9))
    p = txt.paragraphs[0]
    rn = p.add_run(f'{number}  ')
    rn.font.color.rgb = COLOR_PRIMARY
    rn.font.bold = True
    rn.font.size = Pt(15)
    rt = p.add_run(title)
    rt.font.color.rgb = COLOR_TEXT
    rt.font.bold = True
    rt.font.size = Pt(15)
    doc.add_paragraph()

def add_subsection(doc, title):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.font.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = COLOR_PRIMARY

def body(doc, text):
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.size = Pt(11)
        r.font.color.rgb = COLOR_TEXT

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 * (level + 1))
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = COLOR_TEXT

def tip(doc, text, icon='💡', bg=None):
    if bg is None: bg = COLOR_TIP_BG
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    set_cell_border(cell, color_hex=rgb_hex(COLOR_ACCENT), sz=10)
    p = cell.paragraphs[0]
    r = p.add_run(f'{icon}  {text}')
    r.font.size = Pt(10)
    r.font.color.rgb = COLOR_TEXT
    doc.add_paragraph()

def table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_bg(c, COLOR_PRIMARY)
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = COLOR_WHITE
        r.font.size = Pt(10)
    for ri, row_data in enumerate(rows):
        bg = RGBColor(0xF8, 0xF9, 0xFA) if ri % 2 == 0 else COLOR_WHITE
        for ci, val in enumerate(row_data):
            c = t.rows[ri+1].cells[ci]
            set_cell_bg(c, bg)
            r = c.paragraphs[0].add_run(str(val))
            r.font.size = Pt(10)
            r.font.color.rgb = COLOR_TEXT
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()

# ─── BUILD ────────────────────────────────────────────────────────────────────

doc = Document()
for section in doc.sections:
    section.top_margin = section.bottom_margin = Cm(2.5)
    section.left_margin = section.right_margin = Cm(2.8)

# ── COUVERTURE ───────────────────────────────────────────────────────────────
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.add_run('\n\n')
r1 = cover.add_run('NutriVault\n')
r1.font.size = Pt(40); r1.font.bold = True; r1.font.color.rgb = COLOR_PRIMARY
r2 = cover.add_run('Guide d\'utilisation\n')
r2.font.size = Pt(24); r2.font.bold = True; r2.font.color.rgb = COLOR_ACCENT
r3 = cover.add_run('Templates de notes de consultation\n\n')
r3.font.size = Pt(18); r3.font.color.rgb = COLOR_TEXT
r4 = cover.add_run('Version 8.7.16  ·  À l\'attention des diététicien(ne)s\n\n')
r4.font.size = Pt(11); r4.font.color.rgb = COLOR_MUTED

doc.add_page_break()

# ── 1. INTRODUCTION ──────────────────────────────────────────────────────────
add_section_header(doc, '1', 'Introduction')
body(doc, 'Les templates de notes de consultation permettent de standardiser vos prises de notes lors des séances avec vos patients. Préparez votre modèle une seule fois — NutriVault l\'applique automatiquement à chaque nouvelle consultation.')
doc.add_paragraph()
body(doc, 'Ce que vous pouvez inclure dans un template :')
table(doc,
    ['Type d\'élément', 'Description'],
    [['Catégorie de champs', 'Regroupement de champs personnalisés (antécédents, habitudes, évaluation…)'],
     ['Mesure clinique', 'Poids, IMC, tension artérielle, glycémie, etc.'],
     ['Instruction', 'Texte de guidage visible pendant la consultation, avec zone de notes libres']],
    col_widths=[5, 11])
body(doc, 'Avantages :')
bullet(doc, 'Gain de temps : structure préparée à l\'avance')
bullet(doc, 'Cohérence : mêmes informations collectées pour chaque type de consultation')
bullet(doc, 'Traçabilité : notes sauvegardées automatiquement toutes les 3 secondes')
doc.add_paragraph()

# ── 2. ACCÉDER AUX TEMPLATES ─────────────────────────────────────────────────
add_section_header(doc, '2', 'Accéder aux templates')
body(doc, 'Dans le menu de navigation latéral, cliquez sur "Templates de consultation".')
add_screenshot(doc, '01-menu-sidebar.png', 'Menu de navigation — "Templates de consultation"')

# ── 3. LISTE DES TEMPLATES ────────────────────────────────────────────────────
add_section_header(doc, '3', 'Liste des templates')
add_screenshot(doc, '02-liste-templates.png', 'Page liste des templates')
body(doc, 'Pour chaque template : nom avec badge de type coloré, description, nombre d\'éléments, visibilité (Privé / Partagé), statut par défaut (⭐), et boutons d\'action.')
table(doc,
    ['Type', 'Usage conseillé'],
    [['Anamnèse (rouge)', 'Première consultation, collecte de l\'historique patient'],
     ['Évaluation (bleu)', 'Bilan nutritionnel, évaluation clinique'],
     ['Plan alimentaire (vert)', 'Consultation de prescription diététique'],
     ['Suivi (orange)', 'Consultations de suivi régulier'],
     ['Général (violet)', 'Usage polyvalent'],
     ['Personnalisé (sarcelle)', 'Tout autre cas spécifique']],
    col_widths=[5, 11])

doc.add_page_break()

# ── 4. CRÉER UN TEMPLATE ──────────────────────────────────────────────────────
add_section_header(doc, '4', 'Créer un nouveau template')
body(doc, 'Cliquez sur le bouton "Nouveau modele" en haut à droite de la liste.')
add_screenshot(doc, '03-bouton-nouveau-template.png', 'Bouton "Nouveau modele"', width_cm=10)
body(doc, 'L\'éditeur s\'ouvre avec la configuration à gauche et, optionnellement, l\'aperçu à droite.')
add_screenshot(doc, '04-editeur-template-vide.png', 'Éditeur de template — vue initiale')

add_subsection(doc, '4.1  Informations générales')
add_screenshot(doc, '05-section-informations.png', 'Section "Informations du template"')
table(doc,
    ['Champ', 'Obligatoire', 'Description'],
    [['Nom', '✅ Oui', 'Ex. : Bilan nutritionnel initial (200 car. max)'],
     ['Type', 'Non', 'Catégorisation du template'],
     ['Description', 'Non', 'Aide à choisir le bon template en consultation'],
     ['Visibilité', 'Non', 'Privé = vous seul(e) · Partagé = toute l\'équipe'],
     ['Couleur', 'Non', 'Code couleur pour identifier le template visuellement'],
     ['Par défaut', 'Non', 'Pré-sélectionné à chaque nouvelle consultation'],
     ['Tags', 'Non', 'Mots-clés : adulte, diabète, grossesse, sport…']],
    col_widths=[4, 2.5, 9.5])
tip(doc, 'Un seul template peut être "par défaut". En activer un nouveau désactive automatiquement l\'ancien.')

add_subsection(doc, '4.2  Ajouter des éléments')
body(doc, 'Cliquez sur l\'un des trois boutons pour ajouter un élément au template :')
add_screenshot(doc, '06-boutons-ajout-elements.png', 'Boutons d\'ajout d\'éléments', width_cm=12)
bullet(doc, '[ + Catégorie ] — catégorie de champs personnalisés')
bullet(doc, '[ + Mesure ] — mesure clinique (poids, IMC, tension…)')
bullet(doc, '[ + Instruction ] — bloc de texte de guidage avec zone de notes libres')
body(doc, 'Réorganisez les éléments avec les flèches ▲ ▼ ou par glisser-déposer.')
doc.add_paragraph()

add_subsection(doc, '4.3  Catégorie de champs personnalisés')
add_screenshot(doc, '07-modal-categorie.png', 'Modal de sélection d\'une catégorie') if os.path.exists(os.path.join(SCREENSHOTS,'07-modal-categorie.png')) else None
body(doc, 'Chaque catégorie affiche son nom, sa description, le nombre de champs inclus et son niveau :')
bullet(doc, '🔵 Patient — données partagées entre toutes les consultations (ex. : antécédents médicaux)')
bullet(doc, '🟠 Par consultation — données spécifiques à cette séance (ex. : évaluation du jour)')
tip(doc, 'Les champs "Patient" s\'affichent dans toutes les consultations sans re-saisie. Les champs "Par consultation" concernent uniquement la séance en cours.')

add_subsection(doc, '4.4  Mesure clinique')
add_screenshot(doc, '08-modal-mesure.png', 'Modal de sélection d\'une mesure clinique')

add_subsection(doc, '4.5  Instruction')
body(doc, 'Une instruction est un texte que vous rédigez à l\'avance pour vous guider. Elle s\'affiche en encadré jaune lors de la consultation. Vous rédigez le titre et le contenu dans l\'éditeur.')
tip(doc, 'Usage typique : liste de questions à poser au patient, rappel de protocole, check-list d\'éléments à évaluer.')

add_subsection(doc, '4.6  Aperçu et sauvegarde')
body(doc, 'Utilisez le bouton "Aperçu" pour voir le rendu exact du template avant de le sauvegarder.')
add_screenshot(doc, '09-editeur-template-rempli.png', 'Éditeur de template avec éléments configurés')

doc.add_page_break()

# ── 5. UTILISER UN TEMPLATE EN CONSULTATION ───────────────────────────────────
add_section_header(doc, '5', 'Utiliser un template en consultation')

add_subsection(doc, '5.1  Démarrer une consultation')
body(doc, 'Ouvrez la fiche du patient, cliquez sur la visite concernée, puis sur "Démarrer la consultation".')
add_screenshot(doc, '12-detail-visite.png', 'Page de détail d\'une visite')
add_screenshot(doc, '13-bouton-demarrer-consultation.png', 'Bouton "Démarrer la consultation"', width_cm=8)
body(doc, 'Une fenêtre s\'ouvre pour sélectionner le template à utiliser.')
add_screenshot(doc, '14-modal-selection-template.png', 'Modal de sélection du template')

doc.add_page_break()

add_subsection(doc, '5.2  Remplir la note de consultation')
body(doc, 'NutriVault crée automatiquement la note et vous redirige vers l\'éditeur. Les éléments du template s\'affichent dans l\'ordre configuré.')
add_screenshot(doc, '15-editeur-note-vue-ensemble.png', 'Éditeur de note de consultation — vue d\'ensemble')
body(doc, 'Pour chaque catégorie de champs : remplissez les champs normalement (texte, liste déroulante, case à cocher…)')
add_screenshot(doc, '17-carte-categorie-note.png', 'Carte catégorie avec champs à remplir')
body(doc, 'La zone de résumé en bas de page permet de saisir une synthèse globale de la consultation.')
add_screenshot(doc, '22-zone-resume.png', 'Zone de résumé global en bas de page')

add_subsection(doc, '5.3  Enregistrement automatique')
body(doc, 'Vos notes sont sauvegardées automatiquement 3 secondes après chaque modification. L\'en-tête affiche le statut en temps réel.')
add_screenshot(doc, '16-entete-autosave.png', 'En-tête de la note avec indicateurs d\'état')
add_screenshot(doc, '16b-autosave-saving.png', 'Sauvegarde en cours (spinner)', width_cm=12)
add_screenshot(doc, '16c-autosave-saved.png', 'Sauvegarde confirmée (vert)', width_cm=12)
tip(doc, 'Ne fermez pas l\'onglet si l\'indicateur "Sauvegarde en cours" est visible. Le bouton "Enregistrer" permet aussi une sauvegarde manuelle à tout moment.', icon='⚠️', bg=COLOR_WARN_BG)

add_subsection(doc, '5.4  Finaliser et facturer')
body(doc, 'Cliquez sur "Terminer & Facturer" quand la consultation est terminée.')
add_screenshot(doc, '19-bouton-terminer.png', 'Bouton "Terminer & Facturer"', width_cm=8) if os.path.exists(os.path.join(SCREENSHOTS,'19-bouton-terminer.png')) else None
add_screenshot(doc, '20-modal-terminer-facturer.png', 'Modal de finalisation') if os.path.exists(os.path.join(SCREENSHOTS,'20-modal-terminer-facturer.png')) else None
table(doc,
    ['Option', 'Description'],
    [['✅ Marquer la visite comme TERMINÉE', 'Change le statut de la visite dans NutriVault'],
     ['✅ Générer automatiquement une facture', 'Crée une facture associée à cette visite'],
     ['☐ Envoyer la facture par email', 'Envoie la facture au patient (si email renseigné)']],
    col_widths=[7, 9])
tip(doc, 'Une fois confirmée, la note passe en statut "Complétée" et ne peut plus être modifiée.', icon='⚠️', bg=COLOR_WARN_BG)

doc.add_page_break()

# ── 6. GÉRER SES TEMPLATES ────────────────────────────────────────────────────
add_section_header(doc, '6', 'Gérer ses templates')
add_screenshot(doc, '21-liste-templates-actions.png', 'Liste des templates avec boutons d\'action')
body(doc, 'Pour chaque template, trois actions sont disponibles :')
bullet(doc, 'Modifier (icône crayon) — ouvre l\'éditeur avec la configuration actuelle')
bullet(doc, 'Dupliquer — crée une copie identique avec le suffixe (Copie)')
bullet(doc, 'Supprimer — désactive le template si des notes ont déjà été créées avec lui')
doc.add_paragraph()
tip(doc, 'Modifier un template ne modifie PAS les notes déjà créées. Seules les nouvelles consultations bénéficieront des changements.', icon='⚠️', bg=COLOR_WARN_BG)

# ── 7. BONNES PRATIQUES ────────────────────────────────────────────────────────
add_section_header(doc, '7', 'Bonnes pratiques')
add_subsection(doc, 'Organisez par type de consultation')
bullet(doc, '"Bilan initial adulte" → type Anamnèse')
bullet(doc, '"Suivi mensuel" → type Suivi')
bullet(doc, '"Consultation diabète" → type Évaluation')
doc.add_paragraph()
add_subsection(doc, 'Utilisez les tags')
body(doc, 'Ajoutez des tags pour retrouver facilement vos templates : adulte, enfant, diabète, grossesse, sport…')
doc.add_paragraph()
add_subsection(doc, 'Définissez un template par défaut')
body(doc, 'Si vous avez un type de consultation récurrent, marquez-le comme "Par défaut" — il sera pré-sélectionné automatiquement.')
doc.add_paragraph()
add_subsection(doc, 'Partagez avec votre équipe')
body(doc, 'En cabinet collectif, passez vos templates en "Partagé" — toute l\'équipe peut les utiliser.')
doc.add_paragraph()
add_subsection(doc, 'Instructions comme aide-mémoire')
body(doc, 'Rédigez des instructions détaillées pour vos protocoles : questions à poser, éléments à évaluer, grilles de scoring. Vos notes libres documenteront les réponses du patient directement en regard du protocole.')

# ── PIED DE PAGE ──────────────────────────────────────────────────────────────
for section in doc.sections:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('NutriVault · Guide Templates de notes de consultation · v8.7.16')
    r.font.size = Pt(8)
    r.font.color.rgb = COLOR_MUTED

doc.save(OUTPUT)
print(f'✅ Document généré : {OUTPUT}')

# Rapport des screenshots utilisés
shots_used = [
    '01-menu-sidebar.png','02-liste-templates.png','03-bouton-nouveau-template.png',
    '04-editeur-template-vide.png','05-section-informations.png','06-boutons-ajout-elements.png',
    '07-modal-categorie.png','08-modal-mesure.png','09-editeur-template-rempli.png',
    '12-detail-visite.png','13-bouton-demarrer-consultation.png','14-modal-selection-template.png',
    '15-editeur-note-vue-ensemble.png','16-entete-autosave.png','16b-autosave-saving.png',
    '16c-autosave-saved.png','17-carte-categorie-note.png','19-bouton-terminer.png',
    '20-modal-terminer-facturer.png','21-liste-templates-actions.png','22-zone-resume.png'
]
present = [s for s in shots_used if os.path.exists(os.path.join(SCREENSHOTS, s))]
missing = [s for s in shots_used if not os.path.exists(os.path.join(SCREENSHOTS, s))]
print(f'Screenshots : {len(present)}/{len(shots_used)} présents')
if missing:
    print('Manquants :', missing)
